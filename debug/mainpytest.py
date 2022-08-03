import os
import re
from docx import Document
from openpyxl import Workbook,load_workbook

def excel_process():
    path = 'D:\qt_doc\python_test2\debug\quannew.xlsx'
    # path = os.path.join(os.getcwd(),path)
    print(path)
    wb = load_workbook(path)
    print(wb.sheetnames)
    sheet = wb["part1"]
    text=sheet["F3"].value
    sheet["F3"].value="C15"
    print(sheet["F3"].value) #得到混凝土强度等级
    return "C15"

def excel_xieru(temp1,temp2):
    # path = 'D:\qt_doc\python_test2\debug\quannew.xlsx'
    # wb = load_workbook(path)
    # print(temp1)
    return temp2

def printhaha():
    print("hello,world")
    return "hello,world2"

# haha=(23,1,4)
# haha={'zx':'3','qw':'1','io':'4'}

# excel_xieru(haha)
# excel_process()
# printhaha()
# print(os.getcwd())
    # wb.save('newtest2.xlsx')
    # print(type(text))
    # return "haha2"

# if __name__ =="__main__":
#     excel_process()
    # path = 'jisuanshu.docx'
    # # print(path)
    # path = os.path.join(os.getcwd(),path)
    # doc_jisuan = Document(path)

    # # print(len(doc_jisuan.paragraphs))

    # # str1 = re.compile(r"张拉控制应力值σcon=0.7fptk=0.7×(?P<zhi>.*?)张拉控制应力值")
    # i=0
    # # print(doc_jisuan.paragraphs[0].text)
    # # ret = doc_jisuan.paragraphs[0].clear().text
    # # print(ret)
    # for p in doc_jisuan.paragraphs:
    #     p.text = p.text.replace("0.7×'part1-基本参数'!F17=D2","0.7×'part1-基本参数'!F17=7654")
    #     # print(p.text)
    #     # print(str1.findall(p.text))
    #     # if str1.findall(p.text)!= []:
    #     #     textxiugai = doc_jisuan.paragraphs[i]
    #     #     doc_jisuan.paragraphs[i].clear()
    #     #     print(textxiugai)
    #     # #     pass
    #     # i+=1

    # doc_jisuan.save("new.docx")
    # doc_jisuan.close()
    # pass
